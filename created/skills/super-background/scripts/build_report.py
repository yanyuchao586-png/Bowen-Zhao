#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import os
import unicodedata
import uuid
import sys
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_TEMPLATE = ROOT / "assets" / "customer-background-report-template.docx"
DEFAULT_OUTPUT_DIR = ROOT / "outputs"


def configure_stdio() -> None:
    for stream in (sys.stdin, sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8", errors="backslashreplace")
        except Exception:
            pass


def normalize_text(value):
    if isinstance(value, str):
        return unicodedata.normalize("NFC", value)
    if isinstance(value, list):
        return [normalize_text(item) for item in value]
    if isinstance(value, tuple):
        return [normalize_text(item) for item in value]
    if isinstance(value, dict):
        return {str(normalize_text(k)): normalize_text(v) for k, v in value.items()}
    return value


def collect_question_mark_warnings(value, path="$") -> list[str]:
    warnings = []
    if isinstance(value, str):
        if "??" in value:
            warnings.append(path)
    elif isinstance(value, list):
        for index, item in enumerate(value):
            warnings.extend(collect_question_mark_warnings(item, f"{path}[{index}]"))
    elif isinstance(value, dict):
        for key, item in value.items():
            warnings.extend(collect_question_mark_warnings(item, f"{path}.{key}"))
    return warnings


def load_payload(path: str | None) -> dict:
    if path:
        with open(path, "r", encoding="utf-8-sig") as f:
            return normalize_text(json.load(f))
    raw = sys.stdin.read()
    return normalize_text(json.loads(raw.lstrip("\ufeff")))


def has_non_ascii(value: Path) -> bool:
    return any(ord(ch) > 127 for ch in str(value))


def ascii_file_stem(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", str(value))
    chars = []
    for ch in normalized:
        if ch.isascii() and (ch.isalnum() or ch in "-_ "):
            chars.append(ch)
        elif ch.isspace():
            chars.append("_")
        else:
            chars.append("_")
    stem = "".join(chars).strip(" ._").replace(" ", "_")
    while "__" in stem:
        stem = stem.replace("__", "_")
    return stem or "customer-background-report"


def save_docx_encoding_safe(doc: Document, output: Path) -> Path:
    output.parent.mkdir(parents=True, exist_ok=True)
    if has_non_ascii(output.name):
        tmp = output.parent / f"tmp_{uuid.uuid4().hex}.docx"
        doc.save(str(tmp))
        os.replace(str(tmp), str(output))
    else:
        doc.save(str(output))
    return output


def verify_docx_text(output: Path) -> list[str]:
    warnings = []
    try:
        with zipfile.ZipFile(output, "r") as zf:
            xml = zf.read("word/document.xml").decode("utf-8", errors="replace")
    except Exception as exc:
        return [f"docx_verify_failed:{exc}"]
    if "\ufffd" in xml:
        warnings.append("replacement_character_found")
    if "??" in xml:
        warnings.append("question_mark_sequence_found")
    return warnings


def set_font(run, font="Microsoft YaHei", size=None, bold=False, color=None):
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def clear_document_body(doc: Document) -> None:
    body = doc.element.body
    sect_pr = deepcopy(body.sectPr) if body.sectPr is not None else None
    for child in list(body):
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def resolve_paragraph_style(styles, style_name: str | None):
    if not style_name:
        return None
    aliases = {
        "Heading 1": ("Heading1", "Heading 1"),
        "Heading 2": ("Heading2", "Heading 2"),
        "Heading 3": ("Heading3", "Heading 3"),
        "List Paragraph": ("ListParagraph", "List Paragraph"),
    }
    candidates = aliases.get(style_name, (style_name,))
    for candidate in candidates:
        for style in styles:
            try:
                if style.type == 1 and (style.style_id == candidate or style.name == candidate):
                    return style
            except Exception:
                continue
    return None


def set_paragraph_style(paragraph, style_name: str | None) -> None:
    style = resolve_paragraph_style(paragraph.part.document.styles, style_name)
    if style is None:
        return
    try:
        paragraph.style = style
    except Exception:
        pass


def add_text_paragraph(doc, text, *, size=10.5, bold=False, align=None, color=None, style=None):
    p = doc.add_paragraph()
    set_paragraph_style(p, style)
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, color=color)
    if align is not None:
        p.alignment = align
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        set_paragraph_style(p, "List Paragraph")
        r = p.add_run(str(item))
        set_font(r, size=10.2)


def style_cell(cell, header=False):
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)
        for r in p.runs:
            set_font(r, size=9.4, bold=header)
        if header:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def apply_template_table_pr(tbl, table_pr):
    if table_pr is None:
        return
    try:
        existing = tbl._tbl.tblPr
        if existing is not None:
            tbl._tbl.remove(existing)
        tbl._tbl.insert(0, deepcopy(table_pr))
    except Exception:
        pass


def add_two_col_table(doc, rows, table_pr=None):
    tbl = doc.add_table(rows=0, cols=2)
    apply_template_table_pr(tbl, table_pr)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.add_row().cells
    hdr[0].text = "字段"
    hdr[1].text = "内容"
    for c in hdr:
        style_cell(c, header=True)
    for left, right in rows:
        cells = tbl.add_row().cells
        cells[0].text = str(left)
        cells[1].text = str(right)
        style_cell(cells[0])
        style_cell(cells[1])
    for row in tbl.rows:
        row.cells[0].width = Inches(1.7)
        row.cells[1].width = Inches(4.9)
    return tbl


def add_table_from_spec(doc, spec, table_pr=None):
    title = spec.get("title")
    if title:
        add_text_paragraph(doc, title, size=12.5, bold=True, color="1F2937", style="Heading 2")
    rows = spec.get("rows", [])
    add_two_col_table(doc, rows, table_pr=table_pr)


def heading_style_for(text: str) -> str | None:
    text = str(text).strip()
    if not text:
        return None
    h1_prefixes = ("一、", "二、", "三、", "四、", "五、", "六、", "七、", "八、", "九、", "十、")
    if text.startswith(h1_prefixes):
        return "Heading 1"
    if text.startswith("第") or text.startswith("【") or text[:3].replace(".", "").isdigit():
        return "Heading 2"
    if "关键提示" in text or "核心结论" in text or "行动建议" in text:
        return "Heading 2"
    return None


def add_template_heading(doc, text):
    style = heading_style_for(text)
    if style == "Heading 1":
        add_text_paragraph(doc, text, size=15.5, bold=True, color="0F172A", style=style)
    elif style == "Heading 2":
        add_text_paragraph(doc, text, size=12.8, bold=True, color="1F2937", style=style)
    else:
        add_text_paragraph(doc, text, size=12.5, bold=True, color="1F2937")


def build_doc(payload: dict, template: Path, output: Path) -> Path:
    table_pr = None
    if template.exists():
        doc = Document(str(template))
        if doc.tables:
            table_pr = deepcopy(doc.tables[0]._tbl.tblPr)
        clear_document_body(doc)
    else:
        doc = Document()

    styles = doc.styles
    try:
        normal = styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        normal.font.size = Pt(10.5)
    except Exception:
        pass

    title = payload.get("title", "客户超级背调报告")
    subtitle = payload.get("subtitle", "Customer Background Check Report")
    subject = payload.get("subject", "Unknown Subject")

    add_text_paragraph(doc, title, size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color="1F2937")
    add_text_paragraph(doc, subtitle, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color="4B5563")
    add_text_paragraph(doc, subject, size=14.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color="111827")

    if payload.get("category"):
        add_text_paragraph(doc, payload["category"], size=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, color="4B5563")

    warning = payload.get("warning") or "⚠ 本报告含 AI 检索情报；标注「需复核」项请在重大决策前以官方源核验。禁止编造，无据信息标注「暂未找到」。"
    add_text_paragraph(doc, warning, size=9.5, color="6B7280")

    if payload.get("summary_rows"):
        add_two_col_table(doc, payload["summary_rows"], table_pr=table_pr)

    add_text_paragraph(doc, "目录", size=11.5, bold=True, color="1F2937")

    if payload.get("intro_paragraphs") or payload.get("key_points"):
        add_template_heading(doc, "一、背调摘要（Executive Summary）")

    for para in payload.get("intro_paragraphs", []):
        add_text_paragraph(doc, para)

    if payload.get("key_points"):
        add_template_heading(doc, "关键提示")
        add_bullets(doc, payload["key_points"])

    for section in payload.get("sections", []):
        heading = section.get("heading")
        if heading:
            add_template_heading(doc, heading)
        for para in section.get("paragraphs", []):
            add_text_paragraph(doc, para)
        if section.get("bullets"):
            add_bullets(doc, section["bullets"])
        for table_spec in section.get("tables", []):
            add_table_from_spec(doc, table_spec, table_pr=table_pr)

    if payload.get("sources"):
        add_template_heading(doc, "资料来源")
        add_two_col_table(doc, payload["sources"], table_pr=table_pr)

    disclaimer = payload.get("disclaimer")
    if disclaimer:
        add_text_paragraph(doc, disclaimer, size=9.2, color="6B7280")

    save_docx_encoding_safe(doc, output)
    return output


def main():
    configure_stdio()
    parser = argparse.ArgumentParser(description="Build a customer background report docx from JSON.")
    parser.add_argument("--input", help="Path to JSON payload. Reads stdin if omitted.")
    parser.add_argument("--output", help="Output docx path.")
    parser.add_argument("--template", default=str(DEFAULT_TEMPLATE), help="Template docx path.")
    parser.add_argument("--title", help="Optional fallback title.")
    parser.add_argument("--print-plain-path", action="store_true", help="Print only the output path instead of JSON.")
    parser.add_argument("--allow-question-marks", action="store_true", help="Allow payload text that already contains repeated question marks.")
    args = parser.parse_args()

    payload = load_payload(args.input)
    input_warnings = collect_question_mark_warnings(payload)
    if input_warnings and not args.allow_question_marks:
        print(json.dumps({
            "error": "input_encoding_suspect",
            "message": "Input JSON already contains repeated question marks. Regenerate the JSON as UTF-8 before building the Word report.",
            "fields": input_warnings[:20],
        }, ensure_ascii=True))
        sys.exit(2)

    if args.title and not payload.get("title"):
        payload["title"] = args.title

    if not args.output:
        safe_name = payload.get("file_stem") or payload.get("subject") or "customer-background-report"
        safe_name = ascii_file_stem(safe_name)
        args.output = str(DEFAULT_OUTPUT_DIR / f"{safe_name}_customer_background_report.docx")

    output = Path(args.output)
    build_doc(payload, Path(args.template), output)
    warnings = verify_docx_text(output)
    if input_warnings:
        warnings.extend([f"input_question_marks:{field}" for field in input_warnings[:20]])
    if args.print_plain_path:
        print(output)
    else:
        print(json.dumps({"output": str(output), "warnings": warnings}, ensure_ascii=True))


if __name__ == "__main__":
    main()
